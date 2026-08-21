from __future__ import annotations

from io import BytesIO
from pathlib import Path
import shutil

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "sample-docs" / "it-service-desk"
FONT_PATH = Path("C:/Windows/Fonts/arial.ttf")
BOLD_FONT_PATH = Path("C:/Windows/Fonts/arialbd.ttf")

CASES = [
    ("VPN access intermittent", "INC-2401", "Remote Access", "P2", [91, 74, 63, 88]),
    ("Laptop battery replacement", "REQ-2402", "Hardware", "P3", [68, 82, 77, 94]),
    ("MFA enrollment assistance", "REQ-2403", "Identity", "P2", [84, 79, 91, 86]),
    ("Shared drive permissions", "REQ-2404", "Collaboration", "P3", [73, 66, 81, 76]),
    ("Email delivery delay", "INC-2405", "Messaging", "P2", [58, 71, 69, 83]),
    ("Printer queue stuck", "INC-2406", "Workplace", "P3", [88, 92, 80, 86]),
    ("Security update request", "CHG-2407", "Security", "P2", [76, 85, 89, 93]),
    ("New starter account setup", "REQ-2408", "Onboarding", "P3", [95, 87, 90, 96]),
]


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(BOLD_FONT_PATH if bold else FONT_PATH), size)


def draw_wrapped(draw: ImageDraw.ImageDraw, text: str, xy: tuple[int, int], width: int, size: int, fill: str = "#17202a", bold: bool = False) -> int:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textlength(candidate, font=font(size, bold)) <= width:
            current = candidate
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    y = xy[1]
    for line in lines:
        draw.text((xy[0], y), line, font=font(size, bold), fill=fill)
        y += size + 12
    return y


def make_page(case: tuple[str, str, str, str, list[int]], page: int) -> Image.Image:
    subject, ticket, category, priority, values = case
    image = Image.new("RGB", (1600, 2200), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 1600, 150), fill="#123b5d")
    draw.text((80, 42), "NORTHSTAR IT SERVICE DESK", font=font(48, True), fill="white")
    draw.text((80, 190), f"{ticket}  |  {category}  |  Priority {priority}", font=font(34, True), fill="#123b5d")
    draw_wrapped(draw, subject, (80, 260), 1440, 66, bold=True)
    draw.line((80, 390, 1520, 390), fill="#b9c8d5", width=4)
    if page == 1:
        draw.text((80, 450), "SERVICE SUMMARY", font=font(34, True), fill="#123b5d")
        summary = (
            f"Request {ticket} was logged for {subject.lower()}. The service desk confirmed the user impact, "
            "checked the standard support runbook, and routed the work to the responsible resolver group."
        )
        draw_wrapped(draw, summary, (80, 520), 1440, 30)
        rows = [
            ("Opened", "2026-08-20 09:15 UTC"),
            ("Assigned group", f"{category} Operations"),
            ("Current status", "Resolved - monitoring"),
            ("SLA target", "4 business hours"),
        ]
        y = 760
        for label, value in rows:
            draw.rectangle((80, y, 1520, y + 78), outline="#9db2c2", width=2)
            draw.rectangle((80, y, 430, y + 78), fill="#e9f1f6")
            draw.text((105, y + 20), label, font=font(27, True), fill="#123b5d")
            draw.text((470, y + 20), value, font=font(27), fill="#17202a")
            y += 78
        draw.text((80, 1160), "WEEKLY SERVICE METRICS", font=font(34, True), fill="#123b5d")
        chart_left, chart_top, chart_width, chart_height = 120, 1250, 1360, 440
        draw.line((chart_left, chart_top + chart_height, chart_left + chart_width, chart_top + chart_height), fill="#17202a", width=3)
        draw.line((chart_left, chart_top, chart_left, chart_top + chart_height), fill="#17202a", width=3)
        for index, value in enumerate(values):
            x = chart_left + 170 + index * 290
            bar_height = value * 3.2
            draw.rectangle((x, chart_top + chart_height - bar_height, x + 130, chart_top + chart_height), fill="#2b7a78")
            draw.text((x + 35, chart_top + chart_height - bar_height - 48), str(value), font=font(26, True), fill="#17202a")
            draw.text((x - 15, chart_top + chart_height + 24), f"W{index + 1}", font=font(26), fill="#17202a")
        draw.text((120, 1770), "Chart: first-contact resolution percentage by support week", font=font(26), fill="#52616b")
    else:
        draw.text((80, 450), "AGENT WORKSTATION SNAPSHOT", font=font(34, True), fill="#123b5d")
        draw.rectangle((80, 520, 1520, 1320), fill="#202b33", outline="#78909c", width=5)
        draw.rectangle((80, 520, 1520, 600), fill="#33444f")
        draw.ellipse((110, 548, 132, 570), fill="#e36f56")
        draw.ellipse((145, 548, 167, 570), fill="#f0c75e")
        draw.ellipse((180, 548, 202, 570), fill="#6fba84")
        terminal_lines = [
            "$ servicedesk lookup " + ticket,
            "queue: " + category.lower().replace(" ", "-"),
            "impact: single user; workaround available",
            "runbook_check: PASS",
            "identity_check: PASS",
            "resolution_code: STANDARD_CHANGE",
            "status: monitoring",
        ]
        y = 660
        for line in terminal_lines:
            draw.text((140, y), line, font=font(31), fill="#a8e6a3")
            y += 78
        draw.text((80, 1440), "RESOLUTION NOTES", font=font(34, True), fill="#123b5d")
        notes = (
            "Validate the user identity before changing access. Record the affected device or service, "
            "apply the documented remediation, and confirm the user can complete the original task. "
            "Do not include passwords, tokens, or private customer data in the ticket."
        )
        draw_wrapped(draw, notes, (80, 1510), 1440, 30)
        draw.text((80, 1900), "Synthetic screenshot for OCR testing - no real system data", font=font(25), fill="#52616b")
    return image


def image_bytes(image: Image.Image) -> BytesIO:
    stream = BytesIO()
    image.save(stream, format="PNG")
    stream.seek(0)
    return stream


def build_docx(case: tuple[str, str, str, str, list[int]], path: Path, page_images: list[Image.Image]) -> None:
    subject, ticket, category, priority, values = case
    document = Document()
    document.add_heading(f"IT Service Desk Case {ticket}", level=1)
    document.add_paragraph(f"Subject: {subject}")
    document.add_paragraph(f"Category: {category} | Priority: {priority} | Status: Resolved - monitoring")
    document.add_heading("Service summary", level=2)
    document.add_paragraph(
        f"Request {ticket} concerns {subject.lower()}. The service desk confirmed impact, followed the standard support runbook, and recorded the resolution for audit and handoff."
    )
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    for label, value in [("Opened", "2026-08-20 09:15 UTC"), ("Assigned group", f"{category} Operations"), ("SLA target", "4 business hours")]:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.add_heading("Weekly service chart", level=2)
    document.add_picture(image_bytes(page_images[0]), width=Inches(6.4))
    document.add_heading("Agent workstation snapshot", level=2)
    document.add_picture(image_bytes(page_images[1]), width=Inches(6.4))
    document.add_heading("Resolution notes", level=2)
    document.add_paragraph("Validate identity before changing access. Record the affected service, apply the documented remediation, and confirm the original task works. Never record passwords or tokens.")
    document.save(path)


def main() -> None:
    if OUTPUT.exists():
        shutil.rmtree(OUTPUT)
    OUTPUT.mkdir(parents=True)
    for subject, ticket, category, priority, values in CASES:
        case = (subject, ticket, category, priority, values)
        pages = [make_page(case, 1), make_page(case, 2)]
        pdf_path = OUTPUT / f"{ticket}_{subject.replace(' ', '_')}.pdf"
        pages[0].save(pdf_path, format="PDF", resolution=150.0, save_all=True, append_images=pages[1:])
        build_docx(case, OUTPUT / f"{ticket}_{subject.replace(' ', '_')}.docx", pages)
    print(f"Generated {len(CASES)} PDFs and {len(CASES)} DOCX files in {OUTPUT}")


if __name__ == "__main__":
    main()