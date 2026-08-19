from io import BytesIO

from docx import Document
import pytest

from docs_to_markdown import convert_docx, convert_pdf
from docs_to_markdown.pdf_converter import _native_pdf_markdown


def make_docx() -> bytes:
    document = Document()
    document.add_heading("Conversion Proof", level=1)
    document.add_paragraph("This paragraph came from a DOCX file.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Status"
    table.cell(1, 0).text = "MarkItDown"
    table.cell(1, 1).text = "Working"

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _build_pdf(objects: list[bytes]) -> bytes:
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, body in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{number} 0 obj\n".encode("ascii"))
        output.extend(body)
        output.extend(b"\nendobj\n")
    xref = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode("ascii")
    )
    return bytes(output)


def make_pdf() -> bytes:
    content = b"BT /F1 18 Tf 72 720 Td (PDF Conversion Proof) Tj 0 -30 Td /F1 11 Tf (This paragraph came from a PDF file.) Tj ET"
    return _build_pdf(
        [
            b"<< /Type /Catalog /Pages 2 0 R >>",
            b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
            b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
            b"<< /Length " + str(len(content)).encode("ascii") + b" >>\nstream\n" + content + b"\nendstream",
        ]
    )


def make_blank_pdf() -> bytes:
    return _build_pdf(
        [
            b"<< /Type /Catalog /Pages 2 0 R >>",
            b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] >>",
        ]
    )


def test_converts_docx_stream_to_markdown() -> None:
    markdown = convert_docx(make_docx())

    assert "# Conversion Proof" in markdown
    assert "This paragraph came from a DOCX file." in markdown
    assert "| Name | Status |" in markdown
    assert "| MarkItDown | Working |" in markdown


def test_rejects_empty_content() -> None:
    try:
        convert_docx(b"")
    except ValueError as error:
        assert str(error) == "DOCX content cannot be empty"
    else:
        raise AssertionError("Expected empty content to be rejected")


def test_converts_text_pdf_stream_to_markdown() -> None:
    markdown = convert_pdf(make_pdf())

    assert "PDF Conversion Proof" in markdown
    assert "This paragraph came from a PDF file." in markdown


def test_rejects_empty_pdf_content() -> None:
    try:
        convert_pdf(b"")
    except ValueError as error:
        assert str(error) == "PDF content cannot be empty"
    else:
        raise AssertionError("Expected empty PDF content to be rejected")


def test_native_pdf_rejects_malformed_content() -> None:
    with pytest.raises(Exception):
        _native_pdf_markdown(b"%PDF-1.4\nthis is not a valid PDF\n%%EOF")


def test_native_pdf_returns_empty_for_blank_page() -> None:
    assert _native_pdf_markdown(make_blank_pdf()) == ""