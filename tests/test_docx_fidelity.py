import base64
import re
from io import BytesIO

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

from docs_to_markdown import convert_docx
from docs_to_markdown import converter as converter_module


TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def save_document(document: Document) -> bytes:
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_text = OxmlElement("w:t")
    run_text.text = text
    run.append(run_text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def math_run(text: str) -> OxmlElement:
    run = OxmlElement("m:r")
    run_text = OxmlElement("m:t")
    run_text.text = text
    run.append(run_text)
    return run


def inline_fraction() -> OxmlElement:
    math = OxmlElement("m:oMath")
    fraction = OxmlElement("m:f")
    properties = OxmlElement("m:fPr")
    fraction_type = OxmlElement("m:type")
    fraction_type.set(qn("m:val"), "bar")
    properties.append(fraction_type)
    numerator = OxmlElement("m:num")
    numerator.append(math_run("a"))
    denominator = OxmlElement("m:den")
    denominator.append(math_run("b"))
    fraction.extend([properties, numerator, denominator])
    math.append(fraction)
    return math


def block_equation() -> OxmlElement:
    paragraph = OxmlElement("m:oMathPara")
    math = OxmlElement("m:oMath")
    math.append(math_run("E=m"))
    superscript = OxmlElement("m:sSup")
    superscript.append(OxmlElement("m:sSupPr"))
    base = OxmlElement("m:e")
    base.append(math_run("c"))
    exponent = OxmlElement("m:sup")
    exponent.append(math_run("2"))
    superscript.extend([base, exponent])
    math.append(superscript)
    paragraph.append(math)
    return paragraph


def test_preserves_semantic_structure_and_emphasis() -> None:
    document = Document()
    document.add_heading("Document Structure", level=1)
    document.add_heading("Details", level=2)
    paragraph = document.add_paragraph("Plain text with ")
    paragraph.add_run("bold").bold = True
    paragraph.add_run(" and ")
    paragraph.add_run("italic").italic = True
    paragraph.add_run(" emphasis.")

    markdown = convert_docx(save_document(document))

    assert "# Document Structure" in markdown
    assert "## Details" in markdown
    assert "**bold**" in markdown
    assert "*italic*" in markdown


def test_infers_conservative_headings_from_direct_formatting() -> None:
    document = Document()
    title = document.add_paragraph()
    title_run = title.add_run("Visual Title")
    title_run.bold = True
    title_run.font.size = Pt(18)

    section = document.add_paragraph()
    section_run = section.add_run("1. Visual Section")
    section_run.bold = True
    section_run.font.size = Pt(16)
    section_run.font.color.rgb = RGBColor(0x00, 0xB0, 0xF0)

    subsection = document.add_paragraph()
    subsection.add_run("Subsection").bold = True

    label_sentence = document.add_paragraph()
    label_sentence.add_run("For access, use one of the following roles:").bold = True

    list_item = document.add_paragraph(style="List Bullet")
    list_item.add_run("Bold list item").bold = True

    markdown = convert_docx(save_document(document))

    assert "# Visual Title" in markdown
    assert "## 1. Visual Section" in markdown
    assert "### Subsection" in markdown
    assert "### For access" not in markdown
    assert "### Bold list item" not in markdown


def test_preserves_bulleted_and_numbered_lists() -> None:
    document = Document()
    document.add_paragraph("Alpha", style="List Bullet")
    document.add_paragraph("Beta", style="List Bullet")
    document.add_paragraph("First", style="List Number")
    document.add_paragraph("Second", style="List Number")

    markdown = convert_docx(save_document(document))

    assert re.search(r"^\s*[-*+]\s+Alpha\s*$", markdown, re.MULTILINE)
    assert re.search(r"^\s*[-*+]\s+Beta\s*$", markdown, re.MULTILINE)
    assert re.search(r"^\s*1[.)]\s+First\s*$", markdown, re.MULTILINE)
    assert re.search(r"^\s*2[.)]\s+Second\s*$", markdown, re.MULTILINE)


def test_preserves_unicode_and_external_link() -> None:
    document = Document()
    document.add_paragraph("Café — naïve — 東京 — €")
    paragraph = document.add_paragraph("Reference: ")
    add_hyperlink(paragraph, "Project site", "https://example.com/docs")

    markdown = convert_docx(save_document(document))

    assert "Café — naïve — 東京 — €" in markdown
    assert "[Project site](https://example.com/docs)" in markdown


def test_strips_unsafe_external_link_target() -> None:
    document = Document()
    paragraph = document.add_paragraph("Unsafe reference: ")
    add_hyperlink(paragraph, "Do not execute", "javascript:alert(1)")

    markdown = convert_docx(save_document(document))

    assert "Unsafe reference: Do not execute" in markdown
    assert "javascript:" not in markdown


def test_preserves_inline_and_block_equations_as_latex() -> None:
    document = Document()
    inline = document.add_paragraph("Inline fraction before ")
    inline._p.append(inline_fraction())
    inline.add_run(" after inline fraction.")
    document.add_paragraph("Block equation follows:")
    body = document._element.body
    body.insert(len(body) - 1, block_equation())
    document.add_paragraph("Block equation complete.")

    markdown = convert_docx(save_document(document))

    assert "Inline fraction before $\\frac{a}{b}$ after inline fraction." in markdown
    assert "$$E=mc^{2}$$" in markdown


def test_preserves_table_cells() -> None:
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Status"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "Ready"

    markdown = convert_docx(save_document(document))

    assert "| Name | Status |" in markdown
    assert "| Alpha | Ready |" in markdown


def test_preserves_embedded_image_reference() -> None:
    document = Document()
    document.add_paragraph("Image before")
    document.add_picture(BytesIO(TINY_PNG), width=Inches(0.25))
    document.add_paragraph("Image after")

    markdown = convert_docx(save_document(document))

    assert "Image before" in markdown
    assert "Image after" in markdown
    assert "![](data:image/png;base64," in markdown
    assert "base64...)" not in markdown
    assert base64.b64encode(TINY_PNG).decode("ascii") in markdown


def make_captioned_png() -> bytes:
    image = Image.new("RGB", (60, 60), "white")
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def make_legible_png() -> bytes:
    from PIL import ImageDraw

    image = Image.new("RGB", (240, 60), "white")
    ImageDraw.Draw(image).text((10, 20), "Diagram Label", fill="black")
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


class FakeDocxOCR:
    def extract_text(self, image) -> str:
        return "Recognized Diagram Caption"


def test_embeds_ocr_text_in_image_alt_when_provider_available(monkeypatch) -> None:
    monkeypatch.setattr(converter_module, "detect_ocr_adapter", lambda: FakeDocxOCR())
    document = Document()
    document.add_picture(BytesIO(make_captioned_png()), width=Inches(0.5))

    markdown = convert_docx(save_document(document))

    assert "![Image]" in markdown
    assert "[^ocr-1]" in markdown
    assert "[^ocr-1]: OCR text: Recognized Diagram Caption" in markdown


def test_real_ocr_engine_makes_docx_images_searchable() -> None:
    document = Document()
    document.add_picture(BytesIO(make_legible_png()), width=Inches(2))

    markdown = convert_docx(save_document(document))

    assert "Diagram" in markdown


def test_missing_ocr_provider_leaves_docx_images_unaffected(monkeypatch) -> None:
    monkeypatch.setattr(converter_module, "detect_ocr_adapter", lambda: None)
    document = Document()
    document.add_picture(BytesIO(make_captioned_png()), width=Inches(0.5))

    markdown = convert_docx(save_document(document))

    assert "![](data:image/png;base64," in markdown